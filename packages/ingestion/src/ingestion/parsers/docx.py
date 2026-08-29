"""DOCX parser using python-docx with paragraph and table extraction."""

import asyncio
from typing import BinaryIO, List
from uuid import uuid4
from ingestion.parsers.base import DocumentParser, ParsedDocument, Table
from shared.logging import get_logger
from shared.types import DocumentFormat

logger = get_logger(__name__)


class DocxParser(DocumentParser):
    """Parse DOCX documents with text, paragraphs, and tables."""

    @property
    def supported_formats(self) -> List[DocumentFormat]:
        return [DocumentFormat.DOCX]

    async def parse(self, file: BinaryIO, filename: str) -> ParsedDocument:
        return await asyncio.to_thread(self._parse_sync, file, filename)

    def _parse_sync(self, file: BinaryIO, filename: str) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            logger.error("python-docx is not installed. Cannot parse DOCX files.", error=str(e))
            raise ImportError(
                "python-docx is required to parse DOCX files. Please install it with 'pip install python-docx'."
            ) from e

        try:
            doc = DocxDocument(file)
            content_parts: List[str] = []
            tables: List[Table] = []

            # 1. Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content_parts.append(text)

            # 2. Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_cells)

                if table_rows:
                    headers = table_rows[0]
                    data_rows = table_rows[1:] if len(table_rows) > 1 else []
                    tbl_obj = Table(
                        id=str(uuid4()),
                        headers=headers,
                        rows=data_rows,
                        caption=f"Table {table_idx + 1}",
                        metadata={"table_index": table_idx},
                    )
                    tables.append(tbl_obj)
                    content_parts.append(f"\n{tbl_obj.to_markdown()}\n")

            metadata = {
                "format": "docx",
                "filename": filename,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(tables),
            }

            return ParsedDocument(
                content="\n\n".join(content_parts),
                metadata=metadata,
                tables=tables,
            )
        except Exception as e:
            if not isinstance(e, ImportError):
                logger.error("Failed to parse DOCX document", filename=filename, error=str(e))
            raise
