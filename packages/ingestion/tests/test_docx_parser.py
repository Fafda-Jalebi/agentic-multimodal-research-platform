"""Tests for DOCX parser."""

import io
import pytest
from docx import Document
from ingestion.parsers.docx import DocxParser
from shared.types import DocumentFormat


@pytest.mark.asyncio
async def test_docx_parser_with_text_and_tables():
    parser = DocxParser()
    assert DocumentFormat.DOCX in parser.supported_formats

    # Build a real in-memory docx document
    doc = Document()
    doc.add_paragraph("First paragraph of research document.")
    doc.add_paragraph("Second paragraph with findings.")

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Accuracy"
    table.cell(1, 1).text = "95%"

    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    parsed = await parser.parse(docx_bytes, "experiment.docx")

    assert "First paragraph of research document." in parsed.content
    assert "Second paragraph with findings." in parsed.content
    assert len(parsed.tables) == 1
    assert parsed.tables[0].headers == ["Metric", "Value"]
    assert parsed.tables[0].rows == [["Accuracy", "95%"]]
    assert "| Metric | Value |" in parsed.content
    assert "| Accuracy | 95% |" in parsed.content
    assert parsed.metadata["format"] == "docx"
    assert parsed.metadata["paragraph_count"] >= 2
